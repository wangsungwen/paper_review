from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from datetime import date
from pathlib import Path

OUT = Path(__file__).with_name("GitHub_to_HuggingFace_Spaces_部署與維護操作手冊.docx")
NAVY = "173B64"
BLUE = "2E74B5"
LIGHT_BLUE = "E8F1FA"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "64748B"
GOLD = "B7791F"
RED = "9B1C1C"
GREEN = "237A57"


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def set_row_cant_split(row):
    """Keep a troubleshooting/status record together on one page."""
    trPr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    trPr.append(cant_split)


def set_table_widths(table, widths):
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    total = sum(widths)
    tblW.set(qn("w:w"), str(total))
    tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), "120")
    tblInd.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(widths[i]))
            tcW.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_font(run, name="Microsoft JhengHei", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_paragraph(paragraph, fill, border=None):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)
    if border:
        pBdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "6")
        left.set(qn("w:color"), border)
        pBdr.append(left)
        pPr.append(pBdr)


doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = Inches(0.82)
sec.left_margin = sec.right_margin = Inches(0.9)
sec.header_distance = sec.footer_distance = Inches(0.42)

# Compact reference guide token map, with Chinese font override.
normal = doc.styles["Normal"]
normal.font.name = "Microsoft JhengHei"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.22

for style_name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 18, 9),
    ("Heading 2", 13, BLUE, 13, 6),
    ("Heading 3", 11.5, NAVY, 9, 4),
):
    st = doc.styles[style_name]
    st.font.name = "Microsoft JhengHei"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

for style_name in ("List Bullet", "List Number"):
    st = doc.styles[style_name]
    st.font.name = "Microsoft JhengHei"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    st.font.size = Pt(10.5)
    st.paragraph_format.left_indent = Inches(0.38)
    st.paragraph_format.first_line_indent = Inches(-0.19)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.line_spacing = 1.2

code_style = doc.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
code_style.font.name = "Consolas"
code_style._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
code_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
code_style.font.size = Pt(8.2)
code_style.paragraph_format.left_indent = Inches(0.12)
code_style.paragraph_format.right_indent = Inches(0.12)
code_style.paragraph_format.space_before = Pt(3)
code_style.paragraph_format.space_after = Pt(7)
code_style.paragraph_format.line_spacing = 1.05

caption_style = doc.styles.add_style("Manual Caption", WD_STYLE_TYPE.PARAGRAPH)
caption_style.font.name = "Microsoft JhengHei"
caption_style.font.size = Pt(9)
caption_style.font.color.rgb = RGBColor.from_string(MID_GRAY)
caption_style.paragraph_format.space_before = Pt(3)
caption_style.paragraph_format.space_after = Pt(4)


def add_title(text, size=30, color=NAVY, after=8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    set_font(r, size=size, color=color, bold=True)
    return p


def add_para(text="", bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        a, b = text[:len(bold_prefix)], text[len(bold_prefix):]
        set_font(p.add_run(a), bold=True, color=NAVY)
        set_font(p.add_run(b))
    else:
        set_font(p.add_run(text))
    return p


def add_bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    set_font(p.add_run(text))
    return p


def add_num(text):
    p = doc.add_paragraph(style="List Number")
    set_font(p.add_run(text))
    return p


def add_code(code, label=None):
    if label:
        p = doc.add_paragraph(style="Manual Caption")
        set_font(p.add_run(label), size=9, color=MID_GRAY, bold=True)
    p = doc.add_paragraph(style="Code Block")
    shade_paragraph(p, "F5F7FA", "CBD5E1")
    for idx, line in enumerate(code.strip("\n").splitlines()):
        if idx:
            p.add_run("\n")
        r = p.add_run(line)
        set_font(r, name="Consolas", size=8.2, color="1F2937")
    return p


def add_callout(label, text, kind="info"):
    colors = {"info": (LIGHT_BLUE, BLUE), "warn": ("FFF7E6", GOLD), "danger": ("FDECEC", RED), "ok": ("EAF7F1", GREEN)}
    fill, edge = colors[kind]
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.04)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    shade_paragraph(p, fill, edge)
    set_font(p.add_run(f"{label}："), bold=True, color=edge)
    set_font(p.add_run(text), color=NAVY)
    return p


def add_table(headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(h), size=9.2, bold=True, color=NAVY)
    set_repeat_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        set_row_cant_split(table.rows[-1])
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_font(p.add_run(str(value)), size=9)
    set_table_widths(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def h1(text): doc.add_heading(text, level=1)
def h2(text): doc.add_heading(text, level=2)
def h3(text): doc.add_heading(text, level=3)


# Header / footer
header = sec.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(header.add_run("PaperReview｜部署與維護操作手冊"), size=8.5, color=MID_GRAY)
footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(footer.add_run("GitHub → Hugging Face Spaces｜內部維運文件"), size=8.3, color=MID_GRAY)

# Cover
doc.add_paragraph().paragraph_format.space_after = Pt(74)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run("部署與維護操作手冊"), size=12, color=GOLD, bold=True)
add_title("PaperReview", 34, NAVY, 5)
add_title("從 GitHub 部署到 Hugging Face Spaces", 21, BLUE, 14)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run("Streamlit → Gradio｜ZeroGPU｜Git/Xet｜日常更新與故障排除"), size=12, color=MID_GRAY)
doc.add_paragraph().paragraph_format.space_after = Pt(100)
add_table(["文件項目", "內容"], [
    ("適用專案", "wangsungwen/paper_review → wangsongwen/PaperReview"),
    ("目標讀者", "後續維護者、研究助理、系統管理員"),
    ("文件版本", "1.0"),
    ("編製日期", str(date.today())),
    ("主要環境", "Windows PowerShell、Python 3.12、Gradio 6.20、HF Spaces"),
], [1900, 7460])
add_callout("完成標準", "Space 顯示 RUNNING、公開網址回傳 HTTP 200、Gradio 首頁可載入，且程式庫中不含 API Key。", "ok")
navigation_heading = doc.add_heading("文件導覽", level=1)
navigation_heading.paragraph_format.page_break_before = True
add_para("本手冊以可重複操作為目標，從環境準備、GitHub 原始碼取得、介面遷移、Space 建立、推送、驗證，到日常維護與回滾均提供可複製的 PowerShell 指令與含註解程式碼。")
for item in [
    "第 1–3 章：理解架構、準備帳號與取得程式碼",
    "第 4–7 章：把 Streamlit 改為 Gradio，並設定 ZeroGPU/依賴",
    "第 8–10 章：登入、建立 Space、推送與驗證",
    "第 11–14 章：更新、回滾、安全與故障排除",
    "附錄：完整命令速查、交付檢查表與關鍵程式碼",
]: add_bullet(item)
add_callout("重要限制", "Hugging Face 的方案、免費硬體與 ZeroGPU 資格可能調整。若帳號無法選擇 Docker/CPU，就必須使用 Gradio SDK；若也無 ZeroGPU 主機資格，仍可用一般 Gradio CPU Space（依帳號當下可用項目為準）。", "warn")

h1("1. 部署架構與選型")
h2("1.1 最終架構")
add_table(["層級", "技術", "責任"], [
    ("原始碼", "GitHub", "版本管理、協作、正式 main 分支"),
    ("展示與執行", "Hugging Face Gradio Space", "建置依賴、啟動 app.py、提供公開網址"),
    ("介面", "Gradio Blocks", "上傳論文、API 設定、審查流程、結果與下載"),
    ("推論", "Gemini / OpenAI 相容 API", "三輪多代理人論文審查"),
    ("AI 偵測", "ZeroGPU / GPTZero / 模擬", "AI 文字比例與逐句標示"),
    ("狀態隔離", "每次請求的臨時 JSON", "避免不同使用者共用或洩漏 API Key"),
], [1500, 2700, 5160])
h2("1.2 為何不能直接部署原 Streamlit")
add_para("Hugging Face Space 的 SDK 與硬體必須相容。Docker 可執行 Streamlit，但某些帳號無法建立免費 Docker/CPU Space；ZeroGPU 又只支援 Gradio SDK。因此本專案採用「保留核心邏輯、重寫介面層」的方式，將原 app.py 的 Streamlit 控制項改成 Gradio Blocks。")
add_callout("本次實際決策", "舊 Streamlit 介面保留為 streamlit_app.py；新的 app.py 是 Hugging Face Gradio 入口。這樣既能部署，也保留本機舊版參考。", "info")

h1("2. 前置需求與帳號準備")
h2("2.1 必要工具")
for item in [
    "Git 2.40 以上（Windows 可安裝 Git for Windows）",
    "Python 3.10–3.12；本 Space metadata 指定 Python 3.12",
    "Hugging Face CLI：pip install -U huggingface_hub",
    "GitHub 儲存庫讀取權限；私有庫需 GitHub token 或 SSH",
    "Hugging Face 帳號與具有寫入權限的 User Access Token",
]: add_bullet(item)
add_code(r'''# 顯示工具版本；每一行都應成功
git --version
python --version
pip --version
hf --help''', "PowerShell｜確認工具")
h2("2.2 建立 Hugging Face Token")
add_num("登入 Hugging Face 網站，開啟 Settings → Access Tokens。")
add_num("建立 Fine-grained 或 Write token，至少允許管理目標 Space。")
add_num("Token 只貼給 hf auth login，不要放進 README、config.json、程式碼或 Git remote URL。")
add_callout("安全", "任何以 hf_、sk-、AIza 開頭的金鑰，都應視為敏感資料。若曾提交到 Git，即使後來刪除檔案也要立即撤銷並重建金鑰。", "danger")

h1("3. 從 GitHub 取得與檢查原始碼")
h2("3.1 乾淨複製（建議）")
add_code(r'''# 切換到預計放置專案的父資料夾
Set-Location D:\

# 從 GitHub 複製正式儲存庫
git clone https://github.com/wangsungwen/paper_review.git paper_review_system_multi

# 進入專案並確認分支、遠端與工作樹
Set-Location D:\paper_review_system_multi
git branch --show-current
git remote -v
git status --short''', "PowerShell｜Clone 與檢查")
h2("3.2 已有工作區的 worktree 錯誤")
add_para("若 git status 回報 Invalid path '/sessions'，通常是 .git/config 的 core.worktree 指向舊容器路徑。先備份，再修正為目前 Windows 路徑。")
add_code(r'''# 先查看設定；不要直接刪除 .git
Get-Content .git\config

# 正確範例（在 .git/config 的 [core] 內）
# worktree = D:/paper_review_system_multi

# 修正後重新測試
git status --short
git log -1 --oneline''', "PowerShell｜修復 worktree")
add_callout("保護現有修改", "若 git status 顯示大量 M，先用 git diff --numstat 與 git diff -- <檔名> 判斷是內容差異、換行或檔案權限。不要使用 git reset --hard。", "warn")
h2("3.3 .gitignore 必要內容")
add_code(r'''# Python 快取與虛擬環境
__pycache__/
*.py[cod]
.venv/

# 敏感設定與大型模型
config.json
local_models/
*.gguf

# 建置產物與暫存
dist/
build/
.pytest_cache/
*.bundle''', ".gitignore 範例")

h1("4. 將 Streamlit 介面改寫為 Gradio")
h2("4.1 遷移原則")
for item in [
    "不要重寫 core/、services/、llm/、models/；介面只負責收集輸入與呈現輸出。",
    "將 st.session_state 改為 Gradio 元件值、gr.State，或單次函式區域變數。",
    "將 st.file_uploader 改為 gr.File(type='filepath')，再用 adapter 餵給既有檔案解析服務。",
    "將 st.button 邏輯移到 button.click(fn=..., inputs=..., outputs=...)。",
    "長時間三輪審查使用 generator 逐步 yield 狀態，讓使用者看到進度。",
    "API Key 寫入 tempfile，流程完成後在 finally 內刪除。",
]: add_bullet(item)
h2("4.2 保留舊入口")
add_code(r'''# 將原 Streamlit app.py 改名保留
Move-Item -LiteralPath app.py -Destination streamlit_app.py

# 新增 app.py 作為 Gradio Space 入口
# 後續 Hugging Face 會依 README 的 app_file: app.py 啟動''', "PowerShell｜保留舊介面")
h2("4.3 Gradio 基本骨架")
add_code(r'''import gradio as gr

with gr.Blocks(title="多代理人論文審查系統") as demo:
    gr.Markdown("# 🎓 多代理人論文審查系統")
    title = gr.Textbox(label="論文標題")
    content = gr.Textbox(label="論文內容", lines=16)
    run_button = gr.Button("啟動審查", variant="primary")
    result = gr.Markdown()

    # click 將元件值傳給函式，並把回傳值寫到 result
    run_button.click(run_review, [title, content], result)

if __name__ == "__main__":
    # queue 讓長任務排隊；Space 會自動處理公開 host/port
    demo.queue(max_size=20).launch()''', "Python｜最小 Gradio Blocks")
h2("4.4 相容既有檔案解析器")
add_code(r'''class UploadedFileAdapter:
    """把 Gradio filepath 包裝成既有 file_service 所需介面。"""
    def __init__(self, path: str):
        self.name = Path(path).name
        self._handle = open(path, "rb")

    def seek(self, offset: int):
        return self._handle.seek(offset)

    def read(self, size: int = -1):
        return self._handle.read(size)

    def getvalue(self):
        # TXT 解析器需要 getvalue；保留原游標位置
        pos = self._handle.tell()
        self._handle.seek(0)
        data = self._handle.read()
        self._handle.seek(pos)
        return data''', "Python｜TXT/PDF/DOCX 上傳 adapter")
h2("4.5 單次 API Key 隔離")
add_code(r'''def run_review(provider, api_key, model_name, ...):
    config = build_config(provider, api_key, model_name)
    config_path = config_service.write_temp_user_config(config)
    try:
        # 每次請求建立自己的 LLMInterface，不共用其他人的 key
        llm = LLMInterface(config_path=config_path)
        orchestrator = PaperReviewOrchestrator(..., llm=llm)
        ...
    finally:
        # 成功或失敗都刪除臨時設定
        try:
            os.remove(config_path)
        except OSError:
            pass''', "Python｜多使用者安全隔離")
h2("4.6 三輪審查的進度輸出")
add_code(r'''def run_review(...):
    yield "第一輪：獨立審查…", {}, "", "", ""
    round_1 = asyncio.run(orchestrator.run_round_1())

    yield "第二輪：交叉辯論…", {}, format_round(round_1), "", ""
    round_2 = asyncio.run(orchestrator.run_round_2())

    yield "第三輪：最終裁決…", {}, format_round(round_1), format_round(round_2), ""
    round_3 = asyncio.run(orchestrator.run_round_3())

    # 最後一次 yield 同時更新狀態、評分與三輪內容
    yield "完成", orchestrator.review_stats, ..., round_3''', "Python｜Generator 進度更新")

h1("5. ZeroGPU 與 AI 偵測")
h2("5.1 使用 spaces.GPU")
add_code(r'''try:
    import spaces
    # 只在執行 AI 偵測時申請 GPU，最多 120 秒
    gpu_task = spaces.GPU(duration=120)
except (ImportError, RuntimeError):
    # 本機沒有 spaces 套件時仍可測試介面
    def gpu_task(function):
        return function

@gpu_task
def run_ai_detection(text, mode, api_key, api_url):
    from core.ai_detector import AIDetector  # 延遲載入大型套件
    detector = AIDetector(config_path=create_temp_config(...))
    return detector.analyze(text)''', "Python｜ZeroGPU 函式標記")
h2("5.2 延遲載入的重要性")
add_para("不要在 app.py 頂層 import core.ai_detector，否則 Space 一啟動就會載入 torch/transformers，甚至下載模型。延遲 import 可縮短啟動時間，並避免模擬模式也消耗 GPU 與記憶體。")
add_callout("配額", "ZeroGPU 有每日使用配額與排隊機制。一般三輪雲端 LLM 審查不需要 GPU；只有本地 Transformer AI 偵測函式需要 @spaces.GPU。", "info")

h1("6. 設定 README 與 requirements.txt")
h2("6.1 README YAML metadata")
add_code(r'''---
title: PaperReview
emoji: 🎓
colorFrom: blue
colorTo: purple
sdk: gradio
sdk_version: 6.20.0
python_version: "3.12"
app_file: app.py
pinned: false
---''', "README.md｜必須位於檔案最前面")
add_callout("常見錯誤", "若 sdk 誤設為 docker，但硬體仍是 ZeroGPU，Space 會顯示 CONFIG_ERROR：ZeroGPU is only available on Gradio SDK。", "danger")
h2("6.2 精簡 requirements.txt")
add_code(r'''gradio==6.20.0       # UI 與 Space 入口
spaces>=0.42.0       # ZeroGPU 裝飾器
aiohttp>=3.9,<4      # 非同步 HTTP
lxml>=5,<7           # Arxiv XML
pypdf>=6,<7          # PDF 文字擷取
python-docx>=1.1,<2  # DOCX 文字擷取
requests>=2.32,<3    # Gemini/OpenAI/GPTZero HTTP
scikit-learn>=1.5,<2 # TF-IDF 參考文獻挑選
torch==2.9.1         # ZeroGPU 支援版本
transformers>=5,<6   # AI 文字偵測模型''', "requirements.txt｜Space 版核心相依")
add_para("不要把 PyInstaller、Windows-only 套件、llama-cpp-python、CUDA 本機編譯套件全部放入 Space requirements；這些會讓建置變慢或失敗。若本機桌面版仍需要完整相依，可另建 requirements_local.txt。")

h1("7. 本機測試與啟動驗證")
h2("7.1 建立隔離環境")
add_code(r'''python -m venv .venv
.\.venv\Scripts\Activate.ps1
python -m pip install --upgrade pip
python -m pip install -r requirements.txt
python -m pip install pytest''', "PowerShell｜安裝")
h2("7.2 語法、單元測試與 import")
add_code(r'''# 編譯所有主要模組，快速找出語法錯誤
python -m compileall -q app.py core llm models services

# 執行既有測試
python -m pytest -q

# 確認 Gradio Blocks 可建立，不必先開瀏覽器
python -c "import app; print(type(app.demo).__name__)"''', "PowerShell｜測試")
h2("7.3 本機啟動")
add_code(r'''python app.py

# 另一個 PowerShell 視窗進行 HTTP 健康檢查
$resp = Invoke-WebRequest -Uri "http://127.0.0.1:7860/" -UseBasicParsing
$resp.StatusCode  # 預期 200''', "PowerShell｜啟動與健康檢查")

h1("8. Hugging Face 登入與 Git 憑證")
h2("8.1 CLI 登入")
add_code(r'''# 互動式貼上 Token；畫面不應把 token 寫入命令歷史
hf auth login

# 確認登入帳號
hf auth whoami''', "PowerShell｜HF CLI")
h2("8.2 將現有 CLI Token 加到 Git credential")
add_code(r'''# hf CLI 已登入但 git push 仍顯示 Password authentication... 時使用
$token = hf auth token
hf auth login --token $token --add-to-git-credential
Remove-Variable token

# 再確認，但不要執行 hf auth token 並截圖或貼到工單
hf auth whoami''', "PowerShell｜修正 Git Authentication failed")
add_callout("禁止做法", "不要把 token 寫成 https://hf_xxx@huggingface.co/...；它可能出現在 Git config、shell history、日誌或截圖。", "danger")

h1("9. 建立或重建 Hugging Face Space")
h2("9.1 用 Python API 建立 Gradio Space")
add_code(r'''from huggingface_hub import HfApi

api = HfApi()  # 自動讀取 hf auth login 的目前 token
url = api.create_repo(
    repo_id="wangsongwen/PaperReview",
    repo_type="space",
    space_sdk="gradio",
    space_hardware="zero-a10g",  # 僅在帳號可使用 ZeroGPU 時
    private=False,
    exist_ok=True,
)
print(url)''', "Python｜create_space.py")
add_code(r'''python create_space.py''', "PowerShell｜執行")
h2("9.2 刪除與重建的風險")
add_para("刪除 Space 會清除 Space repository、Secrets、Variables、討論與設定。只有在 SDK/硬體卡死且已取得明確授權時才使用 delete_repo。刪除前先記錄 Secrets 名稱與遠端 commit。")
add_code(r'''# 高風險：僅限已備份並取得授權
api.delete_repo(
    repo_id="wangsongwen/PaperReview",
    repo_type="space",
    missing_ok=True,
)''', "Python｜刪除 Space（謹慎）")

h1("10. 推送到 Space")
h2("10.1 新增 Space remote")
add_code(r'''git remote add space https://huggingface.co/spaces/wangsongwen/PaperReview

# 如果 remote 已存在，改用 set-url
git remote set-url space https://huggingface.co/spaces/wangsongwen/PaperReview

git remote -v
git fetch space main''', "PowerShell｜Space remote")
h2("10.2 一般推送")
add_code(r'''# 先確認要推送的內容
git status --short
git diff --check

# 提交正式修改
git add app.py streamlit_app.py README.md requirements.txt llm/interface.py
git commit -m "Deploy native Gradio paper review UI"

# 推送目前 main 到 Space main
git push space main:main''', "PowerShell｜標準部署")
h2("10.3 不改動髒工作樹的乾淨部署快照")
add_para("若本機有大量不相關修改，可用暫存 index 建立單一無父節點部署 commit。這會保留工作樹，不會 reset；並可排除大型 DOCX。")
add_code(r'''git fetch space main
$expected = (git rev-parse space/main).Trim()
$index = Join-Path $env:TEMP "paperreview-hf-index"
Remove-Item $index -Force -ErrorAction SilentlyContinue

# 使用獨立 index，不改目前 .git/index
$env:GIT_INDEX_FILE = $index
git add -A -- . ':(exclude)docs/*.docx'
$tree = (git write-tree).Trim()

# 建立單一部署 commit；不包含 GitHub 歷史中的大型二進位檔
$commit = ('Deploy native Gradio paper review UI' |
    git -c user.name=wangsungwen `
        -c user.email=d11306001@o365.ttu.edu.tw `
        commit-tree $tree).Trim()

Remove-Item Env:GIT_INDEX_FILE
Remove-Item $index -Force -ErrorAction SilentlyContinue

# force-with-lease 僅在遠端仍是剛才 fetch 的版本時覆寫，較安全
git push space "${commit}:refs/heads/main" `
    --force-with-lease="main:$expected"''', "PowerShell｜部署快照（進階）")
h2("10.4 大型二進位檔與 Xet")
add_para("Hugging Face 會拒絕一般 Git 推送大型 DOCX、模型或資料集，訊息為 push was rejected because it contains binary files。可採兩種方式：")
add_bullet("部署不需要：從 Space 快照排除，例如 docs/*.docx、*.gguf、dist/。這是本專案建議方式。")
add_bullet("部署需要：依 Hugging Face Xet 文件安裝 git-xet 並追蹤大型檔；不要用普通 Git 反覆重試。")

h1("11. 部署狀態與線上驗證")
h2("11.1 查詢 Space API")
add_code(r'''$space = Invoke-RestMethod `
  -Uri "https://huggingface.co/api/spaces/wangsongwen/PaperReview"

$space | Select-Object id, sha, lastModified, `
  @{N='stage';E={$_.runtime.stage}}, `
  @{N='hardware';E={$_.runtime.hardware}}, `
  @{N='error';E={$_.runtime.errorMessage}} | Format-List''', "PowerShell｜runtime 狀態")
add_table(["stage", "意義", "處置"], [
    ("BUILDING", "安裝 requirements、建立環境", "等待；過久則查看 Build logs"),
    ("APP_STARTING", "已建置，正在 import/launch", "查看 Runtime logs；通常數十秒"),
    ("RUNNING", "應用正常執行", "繼續做 HTTP 與功能測試"),
    ("CONFIG_ERROR", "README SDK/硬體不相容", "修正 YAML 或硬體"),
    ("BUILD_ERROR", "套件或系統建置失敗", "精簡/固定 requirements 版本"),
    ("RUNTIME_ERROR", "app.py 啟動後崩潰", "看 traceback，先在本機 import app"),
], [1900, 3300, 4160])
h2("11.2 公開網址健康檢查")
add_code(r'''$url = "https://wangsongwen-paperreview.hf.space/"
$resp = Invoke-WebRequest -Uri $url -UseBasicParsing -TimeoutSec 30
$resp.StatusCode              # 預期 200
$resp.Content -match "gradio" # 預期 True''', "PowerShell｜HTTP 200")
h2("11.3 人工驗收")
for item in [
    "首頁、設定 Accordion 與三個主要頁籤可顯示。",
    "上傳 TXT/PDF/DOCX 後顯示解析字數，且內容可編輯。",
    "模擬模式可完成三輪審查，不需 API Key。",
    "Gemini/OpenAI 模式缺少 Key 時顯示可理解錯誤，不洩漏伺服器資訊。",
    "完成後可下載 Markdown 報告。",
    "AI 偵測的模擬模式不下載模型；ZeroGPU 模式會排隊並取得 GPU。",
]: add_bullet(item)

h1("12. 日常更新與維護")
h2("12.1 從 GitHub main 更新 Space")
add_code(r'''git fetch origin main
git switch main
git pull --ff-only origin main

# 執行測試
python -m compileall -q app.py core llm models services
python -m pytest -q

# 推到 Space；若 Space 需要不同 README，可在部署分支維護
git push space main:main''', "PowerShell｜一般更新")
h2("12.2 建議分支策略")
add_table(["分支", "用途", "原則"], [
    ("main", "GitHub 正式程式碼", "通過測試後合併"),
    ("codex/gradio-space", "Space 專用 UI/metadata（可選）", "只放 Gradio/README/精簡相依"),
    ("feature/*", "功能開發", "PR 審查後合併"),
], [2100, 3300, 3960])
h2("12.3 回滾")
add_code(r'''# 查看 Space 遠端紀錄
git fetch space main
git log --oneline space/main -10

# 將某個已知正常 commit 推回 main
$good = "<GOOD_COMMIT_SHA>"
$current = (git rev-parse space/main).Trim()
git push space "${good}:refs/heads/main" `
  --force-with-lease="main:$current"''', "PowerShell｜安全回滾")
add_callout("回滾前", "先保存錯誤 commit SHA 與 Runtime logs。force-with-lease 比 --force 安全，可避免覆蓋別人剛推送的版本。", "warn")

h1("13. 安全與隱私維護")
h2("13.1 部署前金鑰掃描")
add_code(r'''# 確認敏感設定確實被 ignore
git check-ignore -v config.json local_models .venv dist build

# 搜尋常見 token 格式；排除本機 config 與虛擬環境
rg -n --hidden `
  -g '!config.json' -g '!.git/**' -g '!.venv/**' `
  "(AIza[0-9A-Za-z_-]{30,}|hf_[0-9A-Za-z]{20,}|sk-[0-9A-Za-z_-]{20,})" .''', "PowerShell｜Secret scan")
h2("13.2 API Key 處理政策")
for item in [
    "使用 gr.Textbox(type='password')，但仍要避免把值 print 到 logs。",
    "每個請求建立獨立 tempfile；finally 必須刪除。",
    "伺服器共用 config.json 只放空白範例，不放真實 key。",
    "若使用 Space Secrets，程式只以 os.environ 讀取，永遠不回傳給前端。",
    "使用者上傳檔案與產生報告放在暫存空間；Space 重啟後不保證保留。",
]: add_bullet(item)

h1("14. 常見錯誤與解法")
add_table(["錯誤訊息/現象", "根因", "處理方式"], [
    ("Invalid path '/sessions'", ".git/config core.worktree 指向舊容器", "備份後改為目前絕對路徑"),
    ("Password authentication...", "HF CLI token 未加入 Git credential", "hf auth login --add-to-git-credential"),
    ("push rejected: binary files", "DOCX/模型未走 Xet", "排除非必要二進位檔或使用 git-xet"),
    ("ZeroGPU is only available on Gradio SDK", "README 是 docker，硬體是 ZeroGPU", "sdk 改 gradio 並重寫 UI"),
    ("402 Payment Required", "帳號不能建立所選 CPU/Docker 硬體", "改用帳號可用的 Gradio/ZeroGPU，或其他平台"),
    ("CONFIG_ERROR", "README metadata 與 Space 設定衝突", "檢查 sdk、app_file、sdk_version"),
    ("BUILD_ERROR", "相依過多/版本不支援", "精簡 requirements；ZeroGPU 使用支援的 torch"),
    ("APP_STARTING 很久", "頂層下載模型或載入 torch", "延遲 import；模擬模式不載模型"),
    ("Gemini 回應 KeyError choices", "誤用 OpenAI 回應格式解析 Gemini", "解析 candidates[0].content.parts"),
    ("畫面正常但審查失敗", "API Key、模型名、Endpoint 或配額", "先用模擬模式，再逐項驗證雲端 API"),
], [2800, 3060, 3500])
h2("14.1 Gemini 正確回應解析")
add_code(r'''data = response.json()
candidates = data.get("candidates", [])
if not candidates:
    return f"Gemini 沒有候選內容：{data.get('promptFeedback', data)}"

parts = candidates[0].get("content", {}).get("parts", [])
result = "".join(part.get("text", "") for part in parts).strip()
return result or "Gemini 模型回應為空"''', "Python｜Gemini REST candidates 格式")

h1("15. 交付與維護檢查表")
h2("15.1 首次部署前")
for item in [
    "□ GitHub main 可 clone，git status 與 remote 正常",
    "□ README 最前面是 Gradio YAML，app_file 指向 app.py",
    "□ requirements.txt 不含桌面打包與本機 CUDA 非必要套件",
    "□ config.json、模型、虛擬環境、build/dist 均被 ignore",
    "□ app.py 可 import，pytest 全部通過",
    "□ HF token 有目標 Space 寫入權限，且已加入 Git credential",
    "□ 大型二進位檔已排除或使用 Xet",
]: add_bullet(item)
h2("15.2 每次更新後")
for item in [
    "□ API 顯示最新 SHA",
    "□ stage 最終為 RUNNING",
    "□ 公開網址 HTTP 200",
    "□ 檔案上傳、模擬審查、報告下載皆可用",
    "□ Runtime logs 無 Token、使用者論文內容或 traceback 洩漏",
    "□ 記錄本次部署 SHA，保留可回滾版本",
]: add_bullet(item)

h1("附錄 A：完整快速部署命令")
add_code(r'''# 1) 取得程式碼
git clone https://github.com/wangsungwen/paper_review.git paper_review_system_multi
Set-Location .\paper_review_system_multi

# 2) 建立環境、安裝、測試
python -m venv .venv
.\.venv\Scripts\Activate.ps1
python -m pip install -r requirements.txt
python -m pip install pytest
python -m compileall -q app.py core llm models services
python -m pytest -q

# 3) HF 登入並提供 Git 憑證
hf auth login
$token = hf auth token
hf auth login --token $token --add-to-git-credential
Remove-Variable token

# 4) 設定 Space remote 並推送
git remote add space https://huggingface.co/spaces/wangsongwen/PaperReview
git fetch space main
git push space main:main

# 5) 驗證
$space = Invoke-RestMethod -Uri "https://huggingface.co/api/spaces/wangsongwen/PaperReview"
$space.runtime.stage
Invoke-WebRequest -Uri "https://wangsongwen-paperreview.hf.space/" -UseBasicParsing''', "PowerShell｜標準流程（假設 Space 已存在且 main 可直接部署）")

h1("附錄 B：重要網址")
add_table(["用途", "網址"], [
    ("GitHub 原始碼", "https://github.com/wangsungwen/paper_review"),
    ("Hugging Face Space", "https://huggingface.co/spaces/wangsongwen/PaperReview"),
    ("Space 公開應用", "https://wangsongwen-paperreview.hf.space/"),
    ("HF Access Tokens", "https://huggingface.co/settings/tokens"),
    ("Gradio Spaces 文件", "https://huggingface.co/docs/hub/spaces-sdks-gradio"),
    ("ZeroGPU 文件", "https://huggingface.co/docs/hub/spaces-zerogpu"),
    ("Space metadata", "https://huggingface.co/docs/hub/spaces-config-reference"),
    ("Xet 大型檔案", "https://huggingface.co/docs/hub/xet"),
], [2500, 6860])

# Keep headings with the first following paragraph and avoid split code blocks.
for paragraph in doc.paragraphs:
    if paragraph.style.name.startswith("Heading"):
        paragraph.paragraph_format.keep_with_next = True
    if paragraph.style.name == "Code Block":
        paragraph.paragraph_format.keep_together = True

# Core document properties.
doc.core_properties.title = "PaperReview：GitHub 到 Hugging Face Spaces 部署與維護操作手冊"
doc.core_properties.subject = "Gradio、ZeroGPU、Git、部署驗證與日常維護"
doc.core_properties.author = "PaperReview 專案維護團隊"
doc.core_properties.keywords = "Hugging Face Spaces, Gradio, ZeroGPU, GitHub, PaperReview"

doc.save(OUT)
print(OUT)
